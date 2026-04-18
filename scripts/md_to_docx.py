"""
Convert PROJECT-REPORT.md → PROJECT-REPORT.docx
Handles: headings, paragraphs, bullet/numbered lists, code blocks,
         markdown tables, horizontal rules, and a special <<TOC>> block
         rendered as a professional table with leader dots + page numbers.
"""
from __future__ import annotations

import pathlib
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color="AAAAAA", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell_no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _add_shading_to_para(para, fill: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _add_left_border_to_para(para, color="AAAAAA", sz="12"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_bottom_border_to_para(para, color="CCCCCC", sz="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), sz)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_col_width(table, col_idx: int, width_cm: float):
    """Set width of a table column."""
    emu = int(width_cm * 914400 / 2.54)  # cm → EMU → twips-ish; actually docx uses twips
    twips = int(width_cm * 567)          # 1 cm ≈ 567 twips
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        cols = tblGrid.findall(qn("w:gridCol"))
        if col_idx < len(cols):
            cols[col_idx].set(qn("w:w"), str(twips))


# ─────────────────────────────────────────────────────────────────────────────
# Inline markdown parser
# ─────────────────────────────────────────────────────────────────────────────

def _parse_inline(para, text: str, base_bold=False, base_italic=False, base_size: float | None = None):
    """Add runs to *para* honouring **bold**, *italic*, and `code` markers."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            run.bold = base_bold
            run.italic = base_italic
            if base_size:
                run.font.size = Pt(base_size)
        full = m.group(0)
        if full.startswith("**"):
            run = para.add_run(m.group(2))
            run.bold = True
            run.italic = base_italic
        elif full.startswith("*"):
            run = para.add_run(m.group(3))
            run.bold = base_bold
            run.italic = True
        else:  # backtick code
            run = para.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        if base_size and not full.startswith("`"):
            run.font.size = Pt(base_size)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.bold = base_bold
        run.italic = base_italic
        if base_size:
            run.font.size = Pt(base_size)


# ─────────────────────────────────────────────────────────────────────────────
# Style configuration
# ─────────────────────────────────────────────────────────────────────────────

BRAND_BLUE = RGBColor(0x1A, 0x1A, 0x6E)   # deep navy blue for headings
BODY_FONT  = "Times New Roman"
CODE_FONT  = "Courier New"
BODY_SIZE  = 12
CODE_SIZE  = 9.5


def _ensure_style(doc, name, kind=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, kind)


def setup_styles(doc: Document):
    # ── Normal ────────────────────────────────────────────────────────────────
    n = doc.styles["Normal"]
    n.font.name = BODY_FONT
    n.font.size = Pt(BODY_SIZE)
    n.paragraph_format.space_after  = Pt(8)
    n.paragraph_format.space_before = Pt(0)
    n.paragraph_format.line_spacing = Pt(18)
    n.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Heading 1 ─────────────────────────────────────────────────────────────
    h1 = doc.styles["Heading 1"]
    h1.font.name  = BODY_FONT
    h1.font.size  = Pt(16)
    h1.font.bold  = True
    h1.font.color.rgb = BRAND_BLUE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(12)
    h1.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER

    # ── Heading 2 ─────────────────────────────────────────────────────────────
    h2 = doc.styles["Heading 2"]
    h2.font.name  = BODY_FONT
    h2.font.size  = Pt(14)
    h2.font.bold  = True
    h2.font.color.rgb = BRAND_BLUE
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after  = Pt(10)
    h2.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT

    # ── Heading 3 ─────────────────────────────────────────────────────────────
    h3 = doc.styles["Heading 3"]
    h3.font.name  = BODY_FONT
    h3.font.size  = Pt(12)
    h3.font.bold  = True
    h3.font.color.rgb = BRAND_BLUE
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after  = Pt(6)
    h3.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT

    # ── Code Block ────────────────────────────────────────────────────────────
    cs = _ensure_style(doc, "Code Block")
    cs.font.name  = CODE_FONT
    cs.font.size  = Pt(CODE_SIZE)
    cs.paragraph_format.space_before = Pt(0)
    cs.paragraph_format.space_after  = Pt(0)
    cs.paragraph_format.left_indent  = Cm(0.6)
    cs.paragraph_format.line_spacing = Pt(13)
    cs.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT

    # ── Fig Caption ───────────────────────────────────────────────────────────
    fc = _ensure_style(doc, "Fig Caption")
    fc.font.name   = BODY_FONT
    fc.font.size   = Pt(10)
    fc.font.italic = True
    fc.paragraph_format.space_before = Pt(4)
    fc.paragraph_format.space_after  = Pt(14)
    fc.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER

    # ── Bullet Item ───────────────────────────────────────────────────────────
    bi = _ensure_style(doc, "Bullet Item")
    bi.font.name  = BODY_FONT
    bi.font.size  = Pt(BODY_SIZE)
    bi.paragraph_format.space_before  = Pt(2)
    bi.paragraph_format.space_after   = Pt(2)
    bi.paragraph_format.left_indent   = Cm(1.0)
    bi.paragraph_format.first_line_indent = Cm(-0.5)
    bi.paragraph_format.line_spacing  = Pt(18)
    bi.paragraph_format.alignment     = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── TOC Entry (chapter) ───────────────────────────────────────────────────
    toc_ch = _ensure_style(doc, "TOC Chapter")
    toc_ch.font.name  = BODY_FONT
    toc_ch.font.size  = Pt(12)
    toc_ch.font.bold  = True
    toc_ch.paragraph_format.space_before = Pt(6)
    toc_ch.paragraph_format.space_after  = Pt(2)
    toc_ch.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT

    # ── TOC Entry (sub-section) ───────────────────────────────────────────────
    toc_sub = _ensure_style(doc, "TOC Sub")
    toc_sub.font.name   = BODY_FONT
    toc_sub.font.size   = Pt(11)
    toc_sub.font.italic = True
    toc_sub.paragraph_format.space_before    = Pt(1)
    toc_sub.paragraph_format.space_after     = Pt(1)
    toc_sub.paragraph_format.left_indent     = Cm(1.2)
    toc_sub.paragraph_format.alignment       = WD_ALIGN_PARAGRAPH.LEFT


# ─────────────────────────────────────────────────────────────────────────────
# TOC table renderer
# ─────────────────────────────────────────────────────────────────────────────

# TEXT_WIDTH_CM must match (page_width - left_margin - right_margin)
TEXT_WIDTH_CM = 15.92   # A4 with 2.54 cm margins each side


def _add_toc_table(doc: Document, toc_rows: list[str]):
    """
    Render the TOC block as a borderless two-column table.
    Sub-sections (prefixed ~~~) are indented and italic.
    Page numbers are right-aligned.

    Row format in markdown:
      | **1. Abstract** | **6** |       ← chapter (bold in both columns)
      | ~~~3.1 Sub      | 8 |           ← sub-section (indented, italic)
    """
    # Filter out separator row and header
    data = []
    for row in toc_rows:
        row = row.strip()
        if not row.startswith("|"):
            continue
        parts = [c.strip() for c in row.strip("|").split("|")]
        if len(parts) < 2:
            continue
        title_raw = parts[0].strip()
        page_raw  = parts[1].strip()
        # Skip the header row  (CHAPTER / SECTION  |  PAGE)
        if re.match(r"^[-: ]+$", page_raw) or page_raw.upper() == "PAGE":
            continue
        data.append((title_raw, page_raw))

    if not data:
        return

    # ── Create the table ─────────────────────────────────────────────────────
    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"

    # Column widths: title takes ~88%, page ~12%
    title_w_cm = TEXT_WIDTH_CM * 0.88
    page_w_cm  = TEXT_WIDTH_CM * 0.12

    # Set column widths via XML grid
    tbl    = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    else:
        for gc in tblGrid.findall(qn("w:gridCol")):
            tblGrid.remove(gc)
    for w_cm in (title_w_cm, page_w_cm):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w_cm * 567)))
        tblGrid.append(gc)

    # Set overall table width
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(TEXT_WIDTH_CM * 567)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # ── Fill rows ─────────────────────────────────────────────────────────────
    for row_idx, (title_raw, page_raw) in enumerate(data):
        is_sub  = title_raw.startswith("~~~")
        is_bold = "**" in title_raw and not is_sub

        title_clean = title_raw.lstrip("~").strip()
        # Remove ** markers for manual bold control
        title_clean = title_clean.replace("**", "").strip()
        page_clean  = page_raw.replace("**", "").strip()

        row = table.rows[row_idx]

        # ── Title cell ────────────────────────────────────────────────────────
        tc = row.cells[0]
        tc.paragraphs[0].clear()
        p_title = tc.paragraphs[0]

        if is_sub:
            p_title.paragraph_format.left_indent = Cm(1.0)
            p_title.paragraph_format.space_before = Pt(2)
            p_title.paragraph_format.space_after  = Pt(2)

            # Leader dots via tab stop at right edge of title column
            tab_pos_twips = int((title_w_cm - 0.3) * 567)
            pPr = p_title._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:leader"), "dot")
            tab.set(qn("w:pos"), str(tab_pos_twips))
            tabs.append(tab)
            pPr.append(tabs)

            run = p_title.add_run(title_clean)
            run.font.name   = BODY_FONT
            run.font.size   = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # Add tab to trigger leader
            p_title.add_run("\t")
        else:
            p_title.paragraph_format.left_indent  = Cm(0)
            p_title.paragraph_format.space_before = Pt(8)
            p_title.paragraph_format.space_after  = Pt(4)

            tab_pos_twips = int((title_w_cm - 0.3) * 567)
            pPr = p_title._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:leader"), "dot")
            tab.set(qn("w:pos"), str(tab_pos_twips))
            tabs.append(tab)
            pPr.append(tabs)

            run = p_title.add_run(title_clean)
            run.font.name  = BODY_FONT
            run.font.size  = Pt(12)
            run.font.bold  = True
            run.font.color.rgb = BRAND_BLUE
            p_title.add_run("\t")

        _set_cell_no_border(tc)
        # Subtle alternating background
        if not is_sub and row_idx % 2 == 0:
            _set_cell_bg(tc, "F7F7FB")

        # ── Page number cell ──────────────────────────────────────────────────
        pc = row.cells[1]
        pc.paragraphs[0].clear()
        p_page = pc.paragraphs[0]
        p_page.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.RIGHT
        p_page.paragraph_format.space_before = p_title.paragraph_format.space_before
        p_page.paragraph_format.space_after  = p_title.paragraph_format.space_after

        run_pg = p_page.add_run(page_clean)
        run_pg.font.name = BODY_FONT
        run_pg.font.size = Pt(11)
        if not is_sub:
            run_pg.font.bold = True
            run_pg.font.color.rgb = BRAND_BLUE
        else:
            run_pg.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        _set_cell_no_border(pc)
        if not is_sub and row_idx % 2 == 0:
            _set_cell_bg(pc, "F7F7FB")

    # Bottom rule under the whole TOC table
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Code block renderer
# ─────────────────────────────────────────────────────────────────────────────

def _add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph(style="Code Block")
        run = p.add_run(line if line else " ")
        run.font.name = CODE_FONT
        run.font.size = Pt(CODE_SIZE)
        _add_shading_to_para(p, "F5F5F5")
        _add_left_border_to_para(p, color="AAAAAA", sz="12")


# ─────────────────────────────────────────────────────────────────────────────
# Markdown table renderer (for non-TOC tables)
# ─────────────────────────────────────────────────────────────────────────────

def _add_markdown_table(doc: Document, rows: list[str]):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    # Drop separator row
    data = [r for r in parsed
            if not all(re.match(r"^[-:]+$", c.strip()) for c in r if c.strip())]
    # Drop header row that is "CHAPTER / SECTION | PAGE"
    data = [r for r in data
            if not (len(r) >= 1 and r[0].upper() in ("CHAPTER / SECTION", "PAGE"))]

    if not data:
        return

    col_count = max(len(r) for r in data)
    table = doc.add_table(rows=len(data), cols=col_count)
    table.style = "Table Grid"

    for i, row in enumerate(data):
        for j in range(col_count):
            text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            _parse_inline(p, text)
            for run in p.runs:
                run.font.name = BODY_FONT
                run.font.size = Pt(11)
            _set_cell_border(cell)
            if i == 0:
                _set_cell_bg(cell, "E8EAF6")
                for run in p.runs:
                    run.bold = True

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Misc paragraph helpers
# ─────────────────────────────────────────────────────────────────────────────

def _add_page_break(doc: Document):
    from docx.enum.text import WD_BREAK
    para = doc.add_paragraph()
    para.add_run().add_break(WD_BREAK.PAGE)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def _add_hr(doc: Document):
    para = doc.add_paragraph()
    _add_bottom_border_to_para(para)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)


# ─────────────────────────────────────────────────────────────────────────────
# Main parser / converter
# ─────────────────────────────────────────────────────────────────────────────

def convert(md_path: str, docx_path: str):
    doc = Document()

    # A4, 1-inch margins ──────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width   = Cm(21.0)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin   = Cm(2.54)
    section.bottom_margin= Cm(2.54)

    setup_styles(doc)

    text = pathlib.Path(md_path).read_text(encoding="utf-8")
    lines = text.splitlines()

    i = 0
    in_code  = False
    in_toc   = False
    in_table = False
    code_lines:  list[str] = []
    table_lines: list[str] = []
    toc_lines:   list[str] = []

    while i < len(lines):
        line = lines[i]

        # ── TOC sentinel ──────────────────────────────────────────────────────
        if line.strip() == "<<TOC>>":
            in_toc = True
            toc_lines = []
            i += 1
            continue
        if line.strip() == "<<TOC_END>>":
            _add_toc_table(doc, toc_lines)
            in_toc = False
            toc_lines = []
            i += 1
            continue
        if in_toc:
            toc_lines.append(line)
            i += 1
            continue

        # ── Fenced code block ──────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                _add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Markdown table ─────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            peek = lines[i] if i < len(lines) else ""
            if not peek.strip().startswith("|"):
                _add_markdown_table(doc, table_lines)
                in_table = False
                table_lines = []
            continue
        elif in_table:
            _add_markdown_table(doc, table_lines)
            in_table = False
            table_lines = []
            # Fall through — process current line normally

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^-{3,}$|^\*{3,}$", line.strip()):
            _add_hr(doc)
            i += 1
            continue

        # ── Page break ────────────────────────────────────────────────────────
        if line.strip() in (r"\newpage", "\\newpage"):
            _add_page_break(doc)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        hm = re.match(r"^(#{1,4})\s+(.*)", line)
        if hm:
            level = len(hm.group(1))
            text_h = hm.group(2).strip()
            style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 3"}
            p = doc.add_paragraph(style=style_map.get(level, "Heading 2"))
            _parse_inline(p, text_h)
            i += 1
            continue

        # ── Bullet list ────────────────────────────────────────────────────────
        bm = re.match(r"^(\s*)[-*●]\s+(.*)", line)
        if bm:
            depth = len(bm.group(1)) // 2
            p = doc.add_paragraph(style="Bullet Item")
            p.paragraph_format.left_indent       = Cm(1.0 + depth * 0.5)
            p.paragraph_format.first_line_indent  = Cm(-0.5)
            p.add_run("• ").bold = True
            _parse_inline(p, bm.group(2).strip())
            i += 1
            continue

        # ── Numbered list ──────────────────────────────────────────────────────
        nm = re.match(r"^(\s*)(\d+)\.\s+(.*)", line)
        if nm:
            depth = len(nm.group(1)) // 2
            num   = nm.group(2)
            p = doc.add_paragraph(style="Bullet Item")
            p.paragraph_format.left_indent       = Cm(1.0 + depth * 0.5)
            p.paragraph_format.first_line_indent  = Cm(-0.5)
            p.add_run(f"{num}. ")
            _parse_inline(p, nm.group(3).strip())
            i += 1
            continue

        # ── Figure caption (*Fig N:…) ──────────────────────────────────────────
        if re.match(r"^\*Fig\s+\d+", line.strip()):
            p = doc.add_paragraph(style="Fig Caption")
            p.add_run(line.strip().strip("*")).italic = True
            i += 1
            continue

        # ── Image placeholder ![…](…) ──────────────────────────────────────────
        im = re.match(r"!\[([^\]]*)\]\([^)]*\)", line.strip())
        if im:
            p = doc.add_paragraph(style="Fig Caption")
            p.add_run(f"[Figure: {im.group(1)}]").italic = True
            i += 1
            continue

        # ── Empty line ─────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph (collect soft-wrapped lines) ─────────────────────
        para_lines = [line]
        while i + 1 < len(lines):
            nxt = lines[i + 1]
            stop = (
                not nxt.strip()
                or nxt.strip().startswith("#")
                or nxt.strip().startswith("|")
                or nxt.strip().startswith("```")
                or nxt.strip() in (r"\newpage", "---", "***", "<<TOC>>", "<<TOC_END>>")
                or re.match(r"^(\s*)[-*●]\s", nxt)
                or re.match(r"^\s*\d+\.\s", nxt)
                or re.match(r"^\*Fig\s+\d+", nxt.strip())
                or re.match(r"^!\[", nxt.strip())
            )
            if stop:
                break
            para_lines.append(nxt)
            i += 1

        full = " ".join(l.strip() for l in para_lines)

        if re.match(r"^\*Fig\s+\d+", full.strip("*").strip()):
            p = doc.add_paragraph(style="Fig Caption")
            p.add_run(full.strip("*")).italic = True
        else:
            p = doc.add_paragraph(style="Normal")
            _parse_inline(p, full)

        i += 1

    doc.save(docx_path)
    size = pathlib.Path(docx_path).stat().st_size
    print(f"DOCX saved -> {docx_path}")
    print(f"Size: {size:,} bytes  ({size/1024:.1f} KB)")


if __name__ == "__main__":
    convert(
        md_path  = r"c:\Users\psing\Desktop\DevOps Agent\PROJECT-REPORT.md",
        docx_path= r"c:\Users\psing\Desktop\DevOps Agent\PROJECT-REPORT.docx",
    )
